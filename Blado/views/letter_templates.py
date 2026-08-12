"""Blado — Templates de courriers RH pour l'industrie et la métallurgie.

Familles :
  A — Embauchage / Contrats
  B — Attestations et Certificats
  C — Sanctions disciplinaires
  D — Rupture de contrat / Départ
  E — Congés et Absences
  F — Promotion / Avancement
  G — Santé et Sécurité au Travail
  H — CNSS et Déclarations sociales
  I — Convocation / Notifications
  J — Divers RH

Chaque fonction _body_<CODE>() retourne le corps spécifique.
"""
from __future__ import annotations

import os
import re
import sys
from datetime import date

# Fallback entreprise (si aucune en base)
_FALLBACK_NOM = "[Entreprise non configurée]"
_FALLBACK_ADRESSE = "[Adresse]"
_FALLBACK_TEL = "[Tél]"
_FALLBACK_EMAIL = "[Email]"
_FALLBACK_DIRECTEUR = "La Direction"


def _find_logo() -> str:
    """Cherche logo.png dans les emplacements possibles."""
    candidates = [
        os.path.join(os.path.dirname(sys.executable), "logo", "logo.png"),
        os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "..", "logo", "logo.png")),
    ]
    base = getattr(sys, "_MEIPASS", "")
    if base:
        candidates.insert(0, os.path.join(base, "logo", "logo.png"))
    for p in candidates:
        if os.path.exists(p):
            return p
    return candidates[0]

_FALLBACK_NOM = "[Entreprise non configurée]"
_FALLBACK_ADRESSE = "[Adresse]"
_FALLBACK_TEL = "[Tél]"
_FALLBACK_EMAIL = "[Email]"
_FALLBACK_DIRECTEUR = "La Direction"

# ── Staff vide pour les aperçus (toutes les valeurs None → les _s() renvoient [placeholder]) ──
PLACEHOLDER_STAFF: dict[str, None] = {
    "id": None, "full_name": None, "professional_category": None,
    "hire_date": None, "hire_date": None, "base_salary": None,
}

# ── Tokens substituables dans les corps personnalisés ──
# Ordre : le plus long d'abord pour éviter les correspondances partielles.
# Volontairement exclus : [date], [date début], [nom] minuscule — utilisés éditorialement.
_STAFF_TOKENS = (
    ("[Nom complet]", "full_name"),
    ("[Matricule]",   "id"),
    ("[Nom]",         "full_name"),
    ("[Poste]",       "professional_category"),
    ("[poste]",       "professional_category"),
)

def _get_ecole(Service: dict | None) -> dict:
    """Retourne un dict avec les infos établissement depuis le Service ou le fallback."""
    if Service:
        adresse_parts = [Service.get("adress", "")]
        if Service.get("city"):
            adresse_parts.append(Service["city"])
        if Service.get("country"):
            adresse_parts.append(Service["country"])
        return {
            "nom": Service.get("label", _FALLBACK_NOM),
            "adresse": ", ".join(filter(None, adresse_parts)),
            "tel": Service.get("tel_1", _FALLBACK_TEL),
            "email": Service.get("email_1", _FALLBACK_EMAIL),
            "directeur": _FALLBACK_DIRECTEUR,
            "logo": _find_logo(),
        }
    return {
        "nom": _FALLBACK_NOM,
        "adresse": _FALLBACK_ADRESSE,
        "tel": _FALLBACK_TEL,
        "email": _FALLBACK_EMAIL,
        "directeur": _FALLBACK_DIRECTEUR,
        "logo": _find_logo(),
    }


def _s(staff: dict | None, key: str, default: str = "") -> str:
    """Extrait une valeur du dictionnaire staff."""
    if not staff:
        return default
    return str(staff.get(key, default) or default)


def _substitute_staff(text: str, staff: dict | None) -> str:
    """Remplace [Nom], [Poste], [Matricule] par les données réelles du staff.

    Seuls les tokens de la liste _STAFF_TOKENS sont substitués, et uniquement
    quand la valeur correspondante dans staff est truthy.
    """
    if not staff or not text:
        return text
    for token, key in _STAFF_TOKENS:
        val = staff.get(key)
        if val:
            text = text.replace(token, str(val))
    return text


def header(staff: dict | None, objet: str, ref: str,
           today: str = "", Service: dict | None = None) -> str:
    if not today:
        today = date.today().strftime("%d/%m/%Y")
    e = _get_ecole(Service)
    nom = _s(staff, "full_name", "[Nom complet]")
    poste = _s(staff, "professional_category", "[Poste]")
    sid = _s(staff, "id", "[Matricule]")
    return (
        f"{e['nom']}\n"
        f"{e['adresse']}\n"
        f"Tél : {e['tel']}  —  {e['email']}\n"
        f"\n"
        f"Réf : {ref}\n"
        f"Date : {today}\n"
        f"\n"
        f"À l'attention de :\n"
        f"  {nom}\n"
        f"  {poste}\n"
        f"  Matricule : {sid}\n"
        f"\n"
        f"Objet : {objet}\n"
        f"\n"
        f"{'─' * 60}\n"
        f"\n"
    )


def footer(Service: dict | None = None) -> str:
    e = _get_ecole(Service)
    return (
        f"\n{'─' * 60}\n"
        f"\n"
        f"Veuillez agréer, Madame, Monsieur, l'expression de mes salutations distinguées.\n"
        f"\n"
        f"{e['directeur']}\n"
        f"{e['nom']}\n"
        f"\n"
        f"Copie : Dossier de l'employé, Direction\n"
    )


def render_body(staff: dict | None, code: str, template: dict | None = None,
                Service: dict | None = None) -> str:
    """Retourne le corps seul (sans en-tête ni pied) pour aperçu et édition.

    Factorise la logique d'extraction auparavant dupliquée dans generate_docx().
    """
    full = build(staff, code, "", "", Service=Service, template=template)
    lines = full.split("\n")
    body_start = 0
    footer_start = len(lines)
    for i, line in enumerate(lines):
        if line.startswith("Objet :"):
            body_start = i
        if line == "─" * 60 and i > body_start + 1:
            footer_start = i
            break
    body_lines = lines[body_start + 2:]  # skip "Objet :" et ligne vide après
    body_only = "\n".join(body_lines).split(f"\n{'─' * 60}\n")[0].strip()
    return body_only


def build(staff: dict | None, code: str, objet: str, ref: str,
          Service: dict | None = None, template: dict | None = None) -> str:
    global _ECOLE
    _ECOLE = _get_ecole(Service)  # rendu dispo pour toutes les fonctions _body_*
    today = date.today().strftime("%d/%m/%Y")
    h = header(staff, objet, ref, today, Service)
    f = footer(Service)
    body = None

    # 1. Corps personnalisé stocké en base (priorité maximale)
    if template and template.get("body_text"):
        body = _substitute_staff(template["body_text"], staff)

    # 2. Lien vivant vers le code source (copie non modifiée d'un builtin)
    if body is None and template and template.get("source_code"):
        src = template["source_code"]
        if re.fullmatch(r"[A-J][0-9]{2}", src):
            body_func = globals().get(f"_body_{src}")
            if body_func:
                body = body_func(staff, today)

    # 3. Fonction directe _body_{code} (builtin ou code sans template)
    if body is None:
        body_func = globals().get(f"_body_{code}")
        if body_func:
            body = body_func(staff, today)

    # 4. Fallback par famille
    if body is None:
        family = code[0] if code else "F"
        body = _family_fallback(family, staff, today)

    return h + body + f


def _family_fallback(family: str, staff: dict | None, today: str) -> str:
    nom = _s(staff, "full_name", "M./Mme [Nom]")
    poste = _s(staff, "professional_category", "[Poste]")
    sid = _s(staff, "id", "[Matricule]")
    fallbacks = {
        "A": (
            f"Madame, Monsieur,\n\n"
            f"Nous faisons suite à [entretien / décision] et vous informons "
            f"des dispositions relatives à votre contrat de travail.\n\n"
            f"[Détails du contrat : type, dates, salaire, conditions]\n\n"
            f"Nous vous invitons à prendre contact avec le secrétariat pour les formalités."
        ),
        "B": (
            f"Madame, Monsieur,\n\n"
            f"Nous vous informons de la décision suivante concernant votre rémunération "
            f"ou vos avantages professionnels.\n\n"
            f"[Détails : montant, date d'effet, référence]\n\n"
            f"La présente décision prend effet à compter de sa notification."
        ),
        "C": (
            f"Madame, Monsieur,\n\n"
            f"Nous portons à votre connaissance les faits suivants :\n\n"
            f"  • [Description précise des faits]\n\n"
            f"Nous vous informons des suites disciplinaires engagées conformément "
            f"au règlement intérieur et au Code du Travail."
        ),
        "D": (
            f"Madame, Monsieur,\n\n"
            f"Nous faisons suite à notre entretien et vous informons de la fin de "
            f"votre contrat de travail à compter du [date].\n\n"
            f"[Détails : motif, indemnités, préavis]\n\n"
            f"Votre solde de tout compte et vos documents de fin de contrat "
            f"vous seront remis dans les délais légaux."
        ),
        "E": (
            f"Madame, Monsieur,\n\n"
            f"Nous faisons suite à votre demande de congé "
            f"du [date de la demande] pour la période du [date début] au [date fin].\n\n"
            f"Après examen, nous vous informons que [décision].\n\n"
            f"[Motif en cas de refus / Rappel des dates en cas d'accord]"
        ),
        "E": (
            f"Madame, Monsieur,\n\n"
            f"Nous faisons suite à votre [démission / fin de contrat / départ en retraite] "
            f"notifiée le [date].\n\n"
            f"Nous vous informons des modalités de votre départ et des documents "
            f"qui vous seront remis."
        ),
        "F": (
            f"Madame, Monsieur,\n\n"
            f"[Nom],\n\n"
            f"Nous portons à votre connaissance l'information suivante : "
            f"[contenu de la communication].\n\n"
            f"Pour toute question, veuillez vous adresser au secrétariat de direction."
        ),
        "G": (
            f"Madame, Monsieur,\n\n"
            f"Nous faisons suite à votre candidature pour le poste de {poste} "
            f"au sein de {_ECOLE['nom']}.\n\n"
            f"Après examen de votre dossier, nous vous informons que [décision].\n\n"
            f"Nous vous remercions de l'intérêt que vous portez à notre établissement."
        ),
        "H": (
            f"Madame, Monsieur le Directeur,\n\n"
            f"Je soussigné(e), {nom}, matricule n°{sid}, {poste},\n"
            f"ai l'honneur de [objet de la demande].\n\n"
            f"[Détails et motivation de la demande]\n\n"
            f"Je vous prie d'agréer, Madame, Monsieur, "
            f"l'expression de mes salutations distinguées.\n\n"
            f"Signature de l'employé"
        ),
        "I": (
            f"Madame, Monsieur,\n\n"
            f"Dans le cadre du programme du Baccalauréat International (IB), "
            f"nous attestons / vous informons que [contenu spécifique IB].\n\n"
            f"La présente [attestation / notification] est délivrée pour servir "
            f"et valoir ce que de droit."
        ),
        "J": (
            f"Madame, Monsieur le Délégué du Personnel,\n\n"
            f"Dans le cadre du dialogue social au sein de {_ECOLE['nom']}, "
            f"nous vous informons que [contenu].\n\n"
            f"Vous voudrez bien en prendre note et en informer les personnels concernés."
        ),
    }
    return fallbacks.get(family, "[Corps du courrier à compléter]")


# ════════════════════════════════════════════════════════════════════
# Famille A — Embauchage et Contrats
# ════════════════════════════════════════════════════════════════════

def _body_A01(staff: dict, ecole: dict, date: date) -> str:
    """Lettre d'embauche — CDI."""
    nom = _s(staff, "full_name")
    return (
        f"Nous avons le plaisir de vous confirmer votre embauche au sein de "
        f"{ecole['nom']} à compter du {_fmt(date)} en qualité de "
        f"{_s(staff, 'professional_category')}, sous contrat à durée indéterminée.\n\n"
        f"Votre période d'essai est de trois (3) mois, renouvelable une fois.\n"
        f"Votre rémunération mensuelle brute est fixée à "
        f"{_fmt_fcfa(staff.get('base_salary'))}.\n"
        f"Votre lieu de travail est : {ecole['adresse']}.\n\n"
        f"Nous vous souhaitons la bienvenue et pleine réussite dans vos fonctions."
    )

def _body_A02(staff: dict, ecole: dict, date: date) -> str:
    """Lettre d'embauche — CDD."""
    nom = _s(staff, "full_name")
    return (
        f"Nous vous confirmons votre engagement au sein de {ecole['nom']} "
        f"à compter du {_fmt(date)} en qualité de "
        f"{_s(staff, 'professional_category')}, sous contrat à durée déterminée.\n\n"
        f"Ce contrat prendra fin le [date de fin].\n"
        f"Votre rémunération mensuelle brute est de "
        f"{_fmt_fcfa(staff.get('base_salary'))}.\n\n"
        f"Nous vous souhaitons pleine réussite dans votre mission."
    )

def _body_A03(staff: dict, ecole: dict, date: date) -> str:
    """Avenant au contrat de travail."""
    nom = _s(staff, "full_name")
    return (
        f"Par le présent avenant à votre contrat de travail en date du [date contrat], "
        f"nous vous informons de la modification suivante :\n\n"
        f"[Décrire la modification : poste, salaire, horaires, lieu de travail...]\n\n"
        f"Les autres clauses de votre contrat restent inchangées.\n"
        f"Nous vous prions de bien vouloir nous retourner un exemplaire signé."
    )

# ════════════════════════════════════════════════════════════════════
# Famille B — Attestations et Certificats
# ════════════════════════════════════════════════════════════════════

def _body_B01(staff: dict, ecole: dict, date: date) -> str:
    """Attestation de travail."""
    nom = _s(staff, "full_name")
    return (
        f"Nous soussignés, {ecole['titre']}, certifions que "
        f"{_s(staff, 'civility')} {nom}, "
        f"matricule {_s(staff, 'id')}, est employé(e) au sein de "
        f"{ecole['nom']} depuis le {_fmt(staff.get('hire_date') or staff.get('hire_date'))} "
        f"en qualité de {_s(staff, 'professional_category')}.\n\n"
        f"La présente attestation est délivrée à l'intéressé(e) pour servir "
        f"et valoir ce que de droit."
    )

def _body_B02(staff: dict, ecole: dict, date: date) -> str:
    """Certificat de travail — fin de contrat."""
    nom = _s(staff, "full_name")
    return (
        f"Nous certifions que {_s(staff, 'civility')} {nom}, "
        f"a été employé(e) au sein de {ecole['nom']} "
        f"du {_fmt(staff.get('hire_date') or staff.get('hire_date'))} "
        f"au {_fmt(staff.get('departure_date') or date)} "
        f"en qualité de {_s(staff, 'professional_category')}.\n\n"
        f"Durant toute la période de son emploi, {nom} a donné entière satisfaction "
        f"et s'est acquitté(e) de ses fonctions avec professionnalisme et diligence.\n\n"
        f"{nom} est libre de tout engagement envers {ecole['nom']}.\n\n"
        f"Le présent certificat est délivré pour faire valoir ce que de droit."
    )

def _body_B03(staff: dict, ecole: dict, date: date) -> str:
    """Attestation de salaire."""
    nom = _s(staff, "full_name")
    return (
        f"Nous attestons que {_s(staff, 'civility')} {nom} perçoit "
        f"au sein de {ecole['nom']} un salaire mensuel brut de "
        f"{_fmt_fcfa(staff.get('base_salary'))}, soit un salaire net mensuel "
        f"d'environ [montant net] FCFA après déduction des cotisations CNSS "
        f"et de l'impôt sur le revenu.\n\n"
        f"La présente attestation est délivrée à la demande de l'intéressé(e)."
    )

# ════════════════════════════════════════════════════════════════════
# Famille C — Discipline et Sanctions
# ════════════════════════════════════════════════════════════════════

def _body_C01(staff: dict, ecole: dict, date: date) -> str:
    """Avertissement écrit."""
    nom = _s(staff, "full_name")
    return (
        f"Nous avons le regret de vous adresser un avertissement écrit "
        f"pour le motif suivant :\n\n"
        f"[Décrire précisément les faits reprochés, avec dates et circonstances]\n\n"
        f"Cet avertissement sera versé à votre dossier personnel.\n"
        f"Nous vous demandons de prendre les mesures nécessaires pour "
        f"qu'une telle situation ne se reproduise plus.\n\n"
        f"Conformément au Règlement Intérieur et au Code du Travail togolais, "
        f"vous disposez d'un délai de 48 heures pour présenter vos observations écrites."
    )

def _body_C02(staff: dict, ecole: dict, date: date) -> str:
    """Mise à pied conservatoire."""
    nom = _s(staff, "full_name")
    return (
        f"En raison de faits graves portés à notre connaissance, "
        f"nous vous notifions une mise à pied conservatoire à effet immédiat, "
        f"dans l'attente de la décision définitive qui sera prise à l'issue "
        f"de la procédure disciplinaire engagée à votre encontre.\n\n"
        f"Motif : [Décrire les faits]\n\n"
        f"Vous êtes convoqué(e) à un entretien préalable le [date] à [heure] "
        f"au bureau du Directeur Général. Vous pouvez vous faire assister "
        f"par un représentant du personnel."
    )

def _body_C03(staff: dict, ecole: dict, date: date) -> str:
    """Lettre de licenciement pour faute."""
    nom = _s(staff, "full_name")
    return (
        f"Suite à l'entretien préalable du [date] et après examen des faits, "
        f"nous vous notifions votre licenciement pour faute [grave/lourde] "
        f"à compter de ce jour.\n\n"
        f"Motifs : [Détailler les faits constitutifs de la faute]\n\n"
        f"Votre solde de tout compte, votre certificat de travail et votre "
        f"dernier bulletin de paie vous seront remis dans les délais légaux.\n"
        f"Nous vous rappelons vos obligations de confidentialité et de "
        f"non-concurrence si elles figurent dans votre contrat."
    )

# ════════════════════════════════════════════════════════════════════
# Famille D — Rupture de contrat et Départ
# ════════════════════════════════════════════════════════════════════

def _body_D01(staff: dict, ecole: dict, date: date) -> str:
    """Lettre de démission — accusé réception."""
    nom = _s(staff, "full_name")
    return (
        f"Nous accusons réception de votre lettre de démission en date "
        f"du {_fmt(date)}, par laquelle vous nous informez de votre décision "
        f"de quitter vos fonctions de {_s(staff, 'professional_category')} "
        f"au sein de {ecole['nom']}.\n\n"
        f"Nous prenons acte de votre démission. Votre préavis de "
        f"[durée] prendra effet le {_fmt(date)} et s'achèvera le [date fin préavis].\n\n"
        f"Nous vous remercions pour votre contribution et vous souhaitons "
        f"pleine réussite dans vos projets futurs."
    )

def _body_D02(staff: dict, ecole: dict, date: date) -> str:
    """Solde de tout compte."""
    nom = _s(staff, "full_name")
    return (
        f"Nous certifions que le solde de tout compte de "
        f"{_s(staff, 'civility')} {nom}, dont le contrat a pris fin "
        f"le {_fmt(staff.get('departure_date') or date)}, a été arrêté "
        f"et réglé conformément aux dispositions du Code du Travail.\n\n"
        f"Ce solde comprend :\n"
        f"  - Salaire du mois en cours\n"
        f"  - Indemnité compensatrice de congés payés\n"
        f"  - Indemnité de préavis (si applicable)\n"
        f"  - Indemnité de licenciement (si applicable)\n\n"
        f"Le présent reçu pour solde de tout compte est signé par {nom} "
        f"sans réserve."
    )

# ════════════════════════════════════════════════════════════════════
# Famille E — Congés et Absences
# ════════════════════════════════════════════════════════════════════

def _body_E01(staff: dict, ecole: dict, date: date) -> str:
    """Validation de congés annuels."""
    nom = _s(staff, "full_name")
    return (
        f"Nous accusons réception de votre demande de congé annuel du [date demande].\n\n"
        f"Votre congé est accordé pour la période du [date début] au [date fin], "
        f"soit [nb jours] jours ouvrés.\n\n"
        f"Vous voudrez bien prendre vos dispositions pour assurer la passation "
        f"de vos dossiers en cours avant votre départ."
    )

def _body_E02(staff: dict, ecole: dict, date: date) -> str:
    """Refus de congé."""
    nom = _s(staff, "full_name")
    return (
        f"Nous faisons suite à votre demande de congé du [date demande].\n\n"
        f"En raison des nécessités de service liées à [motif : pic de production, "
        f"commande urgente, absence d'un collègue...], nous sommes au regret "
        f"de ne pouvoir y donner une suite favorable pour la période demandée.\n\n"
        f"Nous vous proposons de reporter ce congé à la période du [date]."
    )

# ════════════════════════════════════════════════════════════════════
# Famille F — Promotion et Avancement
# ════════════════════════════════════════════════════════════════════

def _body_F01(staff: dict, ecole: dict, date: date) -> str:
    """Notification de promotion."""
    nom = _s(staff, "full_name")
    return (
        f"Nous avons le plaisir de vous informer de votre promotion au poste de "
        f"{_s(staff, 'professional_category')} à compter du {_fmt(date)}.\n\n"
        f"Cette promotion s'accompagne d'une révision de votre rémunération "
        f"à hauteur de {_fmt_fcfa(staff.get('base_salary'))} brut mensuel.\n\n"
        f"Cette décision témoigne de la reconnaissance de vos compétences "
        f"et de votre engagement au service de {ecole['nom']}.\n"
        f"Nous vous adressons toutes nos félicitations."
    )

def _body_F02(staff: dict, ecole: dict, date: date) -> str:
    """Lettre de félicitations."""
    nom = _s(staff, "full_name")
    return (
        f"Nous tenons à vous exprimer notre vive satisfaction pour "
        f"la qualité de votre travail et votre contribution exceptionnelle "
        f"à [projet/résultat spécifique].\n\n"
        f"Votre professionnalisme et votre engagement font honneur à "
        f"{ecole['nom']} et constituent un exemple pour l'ensemble du personnel.\n\n"
        f"Nous vous adressons nos sincères félicitations."
    )

# ════════════════════════════════════════════════════════════════════
# Famille G — Santé et Sécurité au Travail (Industrie)
# ════════════════════════════════════════════════════════════════════

def _body_G01(staff: dict, ecole: dict, date: date) -> str:
    """Déclaration d'accident de travail."""
    nom = _s(staff, "full_name")
    return (
        f"Nous déclarons par la présente l'accident de travail survenu "
        f"le {_fmt(date)} à [heure] sur le site de [lieu].\n\n"
        f"Employé concerné : {_s(staff, 'civility')} {nom}, "
        f"matricule {_s(staff, 'id')}, poste : {_s(staff, 'professional_category')}.\n\n"
        f"Circonstances de l'accident :\n"
        f"[Décrire précisément les circonstances, témoins éventuels]\n\n"
        f"Nature des blessures : [Description]\n"
        f"Premiers soins administrés : [Description]\n\n"
        f"La présente déclaration est adressée à la CNSS dans les 48 heures "
        f"ouvrables, conformément à l'article 114 du Code de Sécurité Sociale."
    )

def _body_G02(staff: dict, ecole: dict, date: date) -> str:
    """Convocation visite médicale — médecine du travail."""
    nom = _s(staff, "full_name")
    return (
        f"Dans le cadre de la médecine du travail, vous êtes convoqué(e) "
        f"à la visite médicale [périodique / d'embauche / de reprise] "
        f"le [date] à [heure].\n\n"
        f"Lieu : [adresse du centre médical / médecin du travail]\n\n"
        f"Cette visite est obligatoire. Le temps passé est considéré "
        f"comme temps de travail effectif et rémunéré comme tel."
    )

def _body_G03(staff: dict, ecole: dict, date: date) -> str:
    """Attribution Équipements de Protection Individuelle (EPI)."""
    nom = _s(staff, "full_name")
    return (
        f"Vous trouverez ci-joint le bordereau de remise de vos Équipements "
        f"de Protection Individuelle (EPI) obligatoires pour votre poste.\n\n"
        f"Casque de sécurité, chaussures de sécurité, gants, lunettes, "
        f"masque anti-poussière — selon votre fiche de poste.\n\n"
        f"Le port des EPI est OBLIGATOIRE sur le site de production.\n"
        f"Tout manquement expose à des sanctions disciplinaires.\n\n"
        f"Vous êtes responsable du bon entretien de vos EPI. "
        f"Toute détérioration doit être signalée immédiatement à votre supérieur."
    )

# ════════════════════════════════════════════════════════════════════
# Famille H — CNSS et Déclarations sociales (Togo)
# ════════════════════════════════════════════════════════════════════

def _body_H01(staff: dict, ecole: dict, date: date) -> str:
    """Demande d'immatriculation CNSS."""
    nom = _s(staff, "full_name")
    return (
        f"Nous vous informons que votre dossier d'immatriculation à la "
        f"Caisse Nationale de Sécurité Sociale (CNSS) a été constitué "
        f"et transmis à l'agence CNSS compétente.\n\n"
        f"Votre numéro CNSS : {_s(staff, 'cnss_number')}\n\n"
        f"Les cotisations sociales (part salariale 4% et part patronale 16,5%) "
        f"seront prélevées et déclarées trimestriellement conformément "
        f"au Code de Sécurité Sociale togolais."
    )

def _body_H02(staff: dict, ecole: dict, date: date) -> str:
    """Attestation de déclaration CNSS."""
    nom = _s(staff, "full_name")
    return (
        f"Nous attestons que {ecole['nom']} déclare régulièrement "
        f"les salaires de {_s(staff, 'civility')} {nom} auprès de la CNSS "
        f"sous le numéro employeur [N° Employeur CNSS].\n\n"
        f"Les cotisations CNSS sont à jour pour la période "
        f"du [date début] au [date fin].\n\n"
        f"La présente attestation est délivrée à la demande de l'intéressé(e)."
    )

# ════════════════════════════════════════════════════════════════════
# Famille I — Convocations et Notifications
# ════════════════════════════════════════════════════════════════════

def _body_I01(staff: dict, ecole: dict, date: date) -> str:
    """Convocation à un entretien."""
    nom = _s(staff, "full_name")
    return (
        f"Vous êtes convoqué(e) à un entretien le [date] à [heure] "
        f"dans le bureau de [personne].\n\n"
        f"Objet : [Motif de l'entretien]\n\n"
        f"Vous pouvez vous faire assister par un représentant du personnel "
        f"si vous le souhaitez."
    )

def _body_I02(staff: dict, ecole: dict, date: date) -> str:
    """Note de service."""
    return (
        f"NOTE DE SERVICE N° [réf]\n\n"
        f"Date : {_fmt(date)}\n"
        f"Destinataires : Tout le personnel\n"
        f"Objet : [Sujet]\n\n"
        f"[Contenu de la note de service]\n\n"
        f"Le Directeur Général"
    )

# ════════════════════════════════════════════════════════════════════
# Famille J — Divers RH
# ════════════════════════════════════════════════════════════════════

def _body_J01(staff: dict, ecole: dict, date: date) -> str:
    """Lettre de recommandation."""
    nom = _s(staff, "full_name")
    return (
        f"Nous avons le plaisir de recommander {_s(staff, 'civility')} {nom}, "
        f"qui a exercé les fonctions de {_s(staff, 'professional_category')} "
        f"au sein de {ecole['nom']} du {_fmt(staff.get('hire_date') or staff.get('hire_date'))} "
        f"au {_fmt(staff.get('departure_date') or date)}.\n\n"
        f"Durant toute cette période, {nom} a fait preuve de sérieux, "
        f"de ponctualité et d'un excellent esprit d'équipe. "
        f"[Paragraphe personnalisé sur les qualités du candidat]\n\n"
        f"Nous recommandons {nom} sans aucune réserve."
    )

def _body_J02(staff: dict, ecole: dict, date: date) -> str:
    """Modification des conditions de travail."""
    nom = _s(staff, "full_name")
    return (
        f"Nous vous informons de la modification suivante de vos conditions "
        f"de travail, prenant effet le [date] :\n\n"
        f"[Décrire la modification : horaires, poste, lieu...]\n\n"
        f"Cette modification est justifiée par [motif : réorganisation, "
        f"impératifs de production, optimisation...].\n"
        f"Vous disposez d'un délai de 15 jours pour nous faire part "
        f"de vos observations éventuelles."
    )

def _body_J03(staff: dict, ecole: dict, date: date) -> str:
    """Invitation — événement entreprise."""
    return (
        f"Nous avons le plaisir de vous inviter à [événement] "
        f"qui se tiendra le [date] à [heure] au [lieu].\n\n"
        f"[Détails de l'événement : programme, tenue, accompagnants...]\n\n"
        f"Nous comptons sur votre présence pour faire de cet événement "
        f"un moment de convivialité et de renforcement de notre esprit d'équipe."
    )
# ════════════════════════════════════════════════════════════════════


def generate_docx(staff: dict | None, code: str, objet: str, ref: str,
                  Service: dict | None = None, output_path: str = "",
                  template: dict | None = None, body_override: str | None = None) -> str:
    """Génère un fichier .docx avec logo, en-tête formaté et corps du courrier.

    body_override : si fourni (ex. corps édité dans le dialogue de génération),
    utilise ce texte plutôt que build()/render_body().
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    e = _get_ecole(Service)
    today = date.today().strftime("%d/%m/%Y")

    # Corps du courrier
    if body_override is not None:
        body_only = body_override
    else:
        body_only = render_body(staff, code, template=template, Service=Service)

    doc = Document()

    # ── Marges ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # ── Logo en haut à gauche ──
    logo_path = e.get("logo", "")
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.6))

    # ── En-tête établissement ──
    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_nom.add_run(e["nom"])
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1F)

    p_adr = doc.add_paragraph()
    p_adr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_adr.add_run(f"{e['adresse']}  —  Tél : {e['tel']}  —  {e['email']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5F, 0x5F, 0x5F)

    # ── Séparateur ──
    doc.add_paragraph("─" * 70)

    # ── Réf & Date ──
    p_ref = doc.add_paragraph()
    run = p_ref.add_run(f"Réf : {ref}")
    run.font.size = Pt(10)
    p_date = doc.add_paragraph()
    run = p_date.add_run(f"Date : {today}")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Destinataire ──
    nom = _s(staff, "full_name", "[Nom complet]")
    poste = _s(staff, "professional_category", "[Poste]")
    sid = _s(staff, "id", "[Matricule]")
    p_dest = doc.add_paragraph()
    p_dest.add_run("À l'attention de :\n").bold = True
    p_dest.add_run(f"  {nom}\n  {poste}\n  Matricule : {sid}")

    doc.add_paragraph()

    # ── Objet ──
    p_obj = doc.add_paragraph()
    run = p_obj.add_run(f"Objet : {objet}")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ── Corps du courrier — chaque paragraphe = un paragraphe DOCX ──
    for para_text in body_only.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(para_text)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ── Pied ──
    doc.add_paragraph("─" * 70)
    p_close = doc.add_paragraph()
    run = p_close.add_run(
        "Veuillez agréer, Madame, Monsieur, l'expression de mes salutations distinguées.")
    run.font.size = Pt(11)

    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    run = p_sig.add_run(e["directeur"])
    run.bold = True
    run.font.size = Pt(11)
    p_sig_nom = doc.add_paragraph()
    run = p_sig_nom.add_run(e["nom"])
    run.font.size = Pt(11)

    doc.add_paragraph()
    p_copy = doc.add_paragraph()
    run = p_copy.add_run("Copie : Dossier de l'employé, Direction")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5F, 0x5F, 0x5F)

    doc.save(output_path)
    return output_path
